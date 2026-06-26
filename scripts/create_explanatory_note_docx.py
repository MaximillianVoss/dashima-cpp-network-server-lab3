from __future__ import annotations

import csv
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCX_PATH = ROOT / "docs" / "explanatory_note.docx"
UML_PNG = ROOT / "docs" / "uml" / "class_diagram.png"
BENCHMARK_CSV = ROOT / "docs" / "benchmark_results.csv"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    paragraph.paragraph_format.space_after = Pt(0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_geometry(table, widths: list[float]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for row in table.rows:
        for index, width in enumerate(widths):
            row.cells[index].width = Inches(width)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_key_value_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, [2.0, 4.35])
    hdr = table.rows[0].cells
    set_cell_text(hdr[0], "Параметр", True)
    set_cell_text(hdr[1], "Значение", True)
    set_cell_shading(hdr[0], "F2F4F7")
    set_cell_shading(hdr[1], "F2F4F7")
    for key, value in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], key, True)
        set_cell_text(cells[1], value)


def add_three_column_table(doc: Document, headers: tuple[str, str, str], rows: list[tuple[str, str, str]]) -> None:
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    set_table_geometry(table, [1.45, 2.25, 2.65])
    for index, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[index], header, True)
        set_cell_shading(table.rows[0].cells[index], "F2F4F7")
    for row_values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row_values):
            set_cell_text(cells[index], value)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header = section.header.paragraphs[0]
    header.text = "Лабораторная работа 3. Сетевой сервер"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.runs[0].font.size = Pt(9)
    header.runs[0].font.color.rgb = RGBColor(90, 90, 90)


def add_title_page(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(80)
    run = title.add_run("Пояснительная записка")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor.from_string("0B2545")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(18)
    subtitle.add_run("Лабораторная работа № 3: «Иерархия классов»").bold = True

    add_key_value_table(
        doc,
        [
            ("Вариант", "20, «Сетевой сервер»"),
            ("Язык", "C++20"),
            ("Тип прикладной программы", "Диалоговая консольная программа"),
            ("Дополнительные условия", "Динамическая загрузка и плагины не выполнялись"),
            ("Система сборки", "CMake"),
            ("Тестовый фреймворк", "Catch2"),
        ],
    )
    doc.add_page_break()


def add_intro(doc: Document) -> None:
    add_heading(doc, "1. Цель и постановка задачи")
    doc.add_paragraph(
        "Цель работы - спроектировать и реализовать иерархию классов для сетевого сервера, "
        "который хранит таблицу передач и обрабатывает пакеты почты, файла и гипертекста."
    )
    doc.add_paragraph(
        "Сервер имеет сетевое имя, IP-адрес и таблицу передач. Каждый элемент таблицы содержит "
        "указатель на описатель пакета. Для пакетов реализованы операции показа полной информации, "
        "получения типа, адресов отправителя и получателя, а также специальные преобразования "
        "mail -> file и file -> hypertext."
    )

    add_heading(doc, "2. Ход выполнения работы")
    add_numbered(
        doc,
        [
            "Выполнено проектирование UML-диаграммы классов и подготовлены заголовочные файлы.",
            "Реализованы классы внутренней логики: пакеты, сервер, таблица передач и статистика.",
            "Реализован собственный шаблонный контейнер HashTable с аддитивным перемешиванием.",
            "Добавлена диалоговая программа по схеме MVC: model, view, controller.",
            "Разработаны unit-тесты для контейнера и классов логики на Catch2.",
            "Добавлены Doxygen-комментарии для публичных методов классов.",
            "Реализован многопоточный расчет процентного состава пакетов и выполнено сравнение времени.",
        ],
    )


def add_design(doc: Document) -> None:
    add_heading(doc, "3. Проектирование и UML")
    doc.add_paragraph(
        "Проект разделен на компоненты model, view и controller. Классы пользовательского интерфейса "
        "не смешиваются с классами предметной области. Это упрощает тестирование: unit-тесты покрывают "
        "контейнер, пакеты, сервер и статистику без необходимости тестировать консольный ввод."
    )
    if UML_PNG.exists():
        doc.add_picture(str(UML_PNG), width=Inches(6.2))
        caption = doc.add_paragraph("Рисунок 1 - UML-диаграмма классов проекта")
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_heading(doc, "4. Структура проекта")
    add_three_column_table(
        doc,
        ("Каталог", "Назначение", "Содержание"),
        [
            ("include/network_server/model", "Заголовки модели", "Пакеты, сервер, таблица передач, статистика"),
            ("src/model", "Реализация модели", "Методы классов предметной области"),
            ("include/network_server/container", "Шаблонный контейнер", "HashTable и AdditiveHash"),
            ("src/view, src/controller", "Диалоговый интерфейс", "ConsoleView и ServerController"),
            ("tests", "Unit-тесты", "Тесты Catch2 для логики и контейнера"),
            ("docs", "Документация", "UML, результаты бенчмарка, пояснительная записка"),
        ],
    )


def add_implementation(doc: Document) -> None:
    add_heading(doc, "5. Описание реализации")
    add_heading(doc, "5.1 Иерархия пакетов", 2)
    add_three_column_table(
        doc,
        ("Класс", "Ответственность", "Основные операции"),
        [
            ("Packet", "Базовый абстрактный пакет", "type, describe, senderAddress, recipientAddress"),
            ("MailPacket", "Пакет почты", "userName, toFilePacket"),
            ("FilePacket", "Пакет файла", "codeType, infoType, toHypertextPacket"),
            ("HypertextPacket", "Пакет гипертекста", "linkCount, links, describe"),
            ("ServerDescriptor", "Полное описание сервера", "addPacket, deletePacket, selectByRecipientPriority"),
            ("PacketStatistics", "Расчет статистики", "calculateSequential, calculateParallel"),
        ],
    )
    add_heading(doc, "5.2 Шаблонная таблица передач", 2)
    doc.add_paragraph(
        "Таблица передач построена на собственном шаблонном классе HashTable<Key, Value>. "
        "Внутри контейнера используются динамический массив бакетов и однонаправленные цепочки "
        "узлов для разрешения коллизий. STL-контейнеры внутри реализации таблицы не используются."
    )
    doc.add_paragraph(
        "Категория итератора - forward iterator. Она соответствует структуре хранения: элементы "
        "последовательно обходятся по бакетам и цепочкам, но случайный доступ к произвольному "
        "индексу не поддерживается."
    )
    add_heading(doc, "5.3 Диалоговая программа", 2)
    add_bullets(
        doc,
        [
            "добавление пакета в таблицу передач;",
            "поиск и удаление пакета по адресу получателя и типу;",
            "выбор пакета по приоритету hypertext, file, mail;",
            "преобразование mail -> file и file -> hypertext;",
            "показ таблицы передач и списка отправителей;",
            "многопоточный расчет процентного состава пакетов.",
        ],
    )


def read_benchmark_rows() -> list[tuple[str, str, str, str]]:
    rows: list[tuple[str, str, str, str]] = []
    if not BENCHMARK_CSV.exists():
        return rows
    with BENCHMARK_CSV.open(newline="", encoding="utf-8") as source:
        for row in csv.DictReader(source):
            rows.append((row["packet_count"], row["sequential_us"], row["parallel_us"], row["total"]))
    return rows


def add_testing(doc: Document) -> None:
    add_heading(doc, "6. Многопоточность")
    doc.add_paragraph(
        "Процентное соотношение количества пакетов каждого типа реализовано в двух вариантах. "
        "Однопоточная версия проходит по таблице один раз. Многопоточная версия запускает три "
        "задачи: отдельно для почты, файлов и гипертекста. Потоки читают таблицу через const-итератор "
        "и не изменяют общие данные."
    )
    rows = read_benchmark_rows()
    if rows:
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        set_table_geometry(table, [1.5, 1.65, 1.65, 1.55])
        for index, header in enumerate(("Пакетов", "1 поток, мкс", "3 потока, мкс", "Всего")):
            set_cell_text(table.rows[0].cells[index], header, True)
            set_cell_shading(table.rows[0].cells[index], "F2F4F7")
        for values in rows:
            cells = table.add_row().cells
            for index, value in enumerate(values):
                set_cell_text(cells[index], value)

    add_heading(doc, "7. Тестирование и анализ")
    add_three_column_table(
        doc,
        ("Проверка", "Команда/инструмент", "Результат"),
        [
            ("Сборка", "cmake --build build", "Успешно"),
            ("Unit-тесты", "ctest --test-dir build --output-on-failure", "7/7 тестов успешно"),
            ("ASan/UBSan", "build-asan + ctest", "7/7 тестов успешно"),
            ("ThreadSanitizer", "setarch $(uname -m) -R ./build-tsan/network_server_tests", "7 test cases, 31 assertions"),
            ("Диалоговый запуск", "printf '0\\n' | ./build/network_server_dialog", "Штатное завершение"),
        ],
    )
    doc.add_paragraph(
        "Для статической проверки включены предупреждения компилятора -Wall -Wextra -Wpedantic "
        "и добавлен конфигурационный файл .clang-tidy. В текущем WSL-окружении clang-tidy, cppcheck, "
        "valgrind и doxygen не установлены."
    )


def add_conclusion(doc: Document) -> None:
    add_heading(doc, "8. Вывод")
    doc.add_paragraph(
        "В ходе работы разработано приложение на C++20 для варианта «Сетевой сервер». "
        "Реализованы иерархия классов, собственная шаблонная таблица передач, диалоговый "
        "интерфейс, модульные тесты, UML-диаграмма и многопоточный расчет статистики. "
        "Дополнительные задания с динамической загрузкой и плагинами не выполнялись."
    )


def main() -> None:
    doc = Document()
    configure_document(doc)
    add_title_page(doc)
    add_intro(doc)
    add_design(doc)
    add_implementation(doc)
    add_testing(doc)
    add_conclusion(doc)
    DOCX_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(DOCX_PATH)


if __name__ == "__main__":
    main()
